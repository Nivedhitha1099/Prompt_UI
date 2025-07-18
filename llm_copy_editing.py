import streamlit as st
import os
import io
import tempfile
from pathlib import Path
import docx
from bs4 import BeautifulSoup
import markdown
import difflib
import re
from langchain_anthropic import ChatAnthropic
from langchain.schema import HumanMessage
from dotenv import load_dotenv
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# Load environment variables
load_dotenv()

# Page configuration
st.set_page_config(
    page_title="Document Copyediting Tool",
    page_icon="📝",
    layout="wide"
)

# Initialize the chat model
@st.cache_resource
def initialize_chat_model():
    try:
        # Fix: Use os.getenv() as a function, not as a dictionary
        token = os.getenv("LLMFOUNDRY_TOKEN")
        if not token:
            raise Exception("LLMFOUNDRY_TOKEN environment variable not found")
            
        chat_model = ChatAnthropic(
            anthropic_api_key=f'{token}:my-test-project',
            anthropic_api_url="https://llmfoundry.straive.com/anthropic/",
            model_name="claude-3-haiku-20240307"
        )
        return chat_model
    except Exception as e:
        st.error(f"Failed to initialize AI model: {str(e)}")
        return None

# Style guide rules
STYLE_GUIDE_RULES = {
    "American English Conventions": {
        "Convert British spellings to American": True,
        "Use American grammar conventions": True,
        "Use American vocabulary": True,
        "Use American quotation mark style": True
    },
    "Capitalization": {
        "Capitalize proper nouns": True,
        "Lowercase articles unless at sentence start": True,
        "Lowercase coordinating conjunctions": True,
        "Convert written years to numerals": True,
        "Capitalize words in organization names": True
    },
    "Grammar and Usage": {
        "Use contractions appropriately": True,
        "Correct misused homophones": True,
        "Correct possessives": True,
        "Remove apostrophes from possessive pronouns": True
    },
    "Punctuation": {
        "Use serial comma": True,
        "Comma before 'etc.' but not after": True,
        "Commas around 'too' in middle, not at end": True,
        "Lowercase first word after colon": True
    },
    "Number Formatting": {
        "Spell out numbers under 10": True,
        "Use numerals for numbers 10 and above": True,
        "Write numbers >999,999 as cardinal numbers": True,
        "Format numbers ≥1,000 with commas": True
    },
    "Dash Usage": {
        "Use hyphens for compound words": True,
        "Use en dashes for ranges": True,
        "Use em dashes for breaks in thought": True
    },
    "List Formatting": {
        "Use proper numbering for ordered lists": True,
        "Use consistent bullets for unordered lists": True,
        "Proper capitalization and punctuation": True
    }
}

def read_txt_file(file):
    """Read text file"""
    try:
        return file.getvalue().decode('utf-8')
    except UnicodeDecodeError:
        return file.getvalue().decode('latin-1')

def read_docx_file(file):
    """Read DOCX file"""
    try:
        doc = docx.Document(file)
        content = []
        for paragraph in doc.paragraphs:
            content.append(paragraph.text)
        return '\n'.join(content)
    except Exception as e:
        raise Exception(f"Error reading DOCX file: {str(e)}")

def read_html_file(file):
    """Read HTML file and extract text"""
    try:
        content = file.getvalue().decode('utf-8')
        soup = BeautifulSoup(content, 'html.parser')
        return soup.get_text()
    except Exception as e:
        raise Exception(f"Error reading HTML file: {str(e)}")

def read_markdown_file(file):
    """Read Markdown file"""
    try:
        content = file.getvalue().decode('utf-8')
        # Convert markdown to plain text for processing
        html = markdown.markdown(content)
        soup = BeautifulSoup(html, 'html.parser')
        return soup.get_text()
    except Exception as e:
        raise Exception(f"Error reading Markdown file: {str(e)}")

def process_document(file):
    """Process uploaded document based on file type"""
    file_extension = Path(file.name).suffix.lower()
    
    try:
        if file_extension == '.txt':
            return read_txt_file(file)
        elif file_extension == '.docx':
            return read_docx_file(file)
        elif file_extension in ['.html', '.htm']:
            return read_html_file(file)
        elif file_extension in ['.md', '.markdown']:
            return read_markdown_file(file)
        else:
            raise Exception(f"Unsupported file type: {file_extension}")
    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
        return None

def create_editing_prompt(text, selected_rules, custom_instructions=""):
    """Create prompt for the AI model"""
    rules_text = []
    for category, rules in selected_rules.items():
        if any(rules.values()):
            active_rules = [rule for rule, active in rules.items() if active]
            rules_text.append(f"{category}: {', '.join(active_rules)}")
    
    # Build the base prompt
    prompt_parts = ["Please copyedit the following text according to these guidelines:"]
    
    # Add style guide rules if any are selected
    if rules_text:
        prompt_parts.append("\nSTYLE GUIDE RULES:")
        prompt_parts.extend([f"- {rule}" for rule in rules_text])
    
    # Add custom instructions if provided
    if custom_instructions.strip():
        prompt_parts.append(f"\nADDITIONAL INSTRUCTIONS:")
        prompt_parts.append(f"{custom_instructions.strip()}")
    
    # Add the original text and final instructions
    prompt_parts.extend([
        f"\nOriginal Text:",
        f"{text}",
        f"\nInstructions:",
        f"1. Apply all the selected copyediting rules and follow any additional instructions",
        f"2. Maintain the original meaning and structure unless otherwise specified",
        f"3. Make necessary corrections but preserve formatting",
        f"4. Return only the corrected text without explanations",
        f"\nCorrected Text:"
    ])
    
    return "\n".join(prompt_parts)

def highlight_differences(original, edited):
    """Create highlighted HTML showing differences between original and edited text"""
    # Split into words for better granular comparison
    original_words = re.findall(r'\S+|\s+', original)
    edited_words = re.findall(r'\S+|\s+', edited)
    
    # Create diff
    differ = difflib.SequenceMatcher(None, original_words, edited_words)
    
    highlighted_original = []
    highlighted_edited = []
    
    for tag, i1, i2, j1, j2 in differ.get_opcodes():
        if tag == 'equal':
            # No change
            original_chunk = ''.join(original_words[i1:i2])
            edited_chunk = ''.join(edited_words[j1:j2])
            highlighted_original.append(original_chunk)
            highlighted_edited.append(edited_chunk)
        elif tag == 'delete':
            # Text removed in edited version
            original_chunk = ''.join(original_words[i1:i2])
            highlighted_original.append(f'<span style="background-color: #ffcccc; text-decoration: line-through;">{original_chunk}</span>')
        elif tag == 'insert':
            # Text added in edited version
            edited_chunk = ''.join(edited_words[j1:j2])
            highlighted_edited.append(f'<span style="background-color: #ccffcc; font-weight: bold;">{edited_chunk}</span>')
        elif tag == 'replace':
            # Text changed
            original_chunk = ''.join(original_words[i1:i2])
            edited_chunk = ''.join(edited_words[j1:j2])
            highlighted_original.append(f'<span style="background-color: #ffcccc; text-decoration: line-through;">{original_chunk}</span>')
            highlighted_edited.append(f'<span style="background-color: #ccffcc; font-weight: bold;">{edited_chunk}</span>')
    
    return ''.join(highlighted_original), ''.join(highlighted_edited)

def format_diff_to_docx(original_docx_file_obj, edited_text, output_path):
    """
    Creates a DOCX document with tracked changes.
    Red strikethrough for deletions, green underline for insertions.
    Attempts to preserve original formatting for equal and deleted text.
    For inserted text, it tries to inherit paragraph and run formatting from nearby original content.
    """
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.oxml.ns import qn
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    # Load the original document from the file object
    # Ensure the file object's pointer is at the beginning
    original_docx_file_obj.seek(0) 
    original_doc = Document(original_docx_file_obj)
    doc = Document()

    doc.add_heading("Tracked Changes", level=1)

    original_paragraphs = [p.text for p in original_doc.paragraphs]
    edited_paragraphs = edited_text.split('\n')

    differ = difflib.SequenceMatcher(None, original_paragraphs, edited_paragraphs)

    for tag, i1, i2, j1, j2 in differ.get_opcodes():
        if tag == 'equal':
            # Copy paragraphs from original document preserving styles
            for idx in range(i1, i2):
                orig_p = original_doc.paragraphs[idx]
                new_p = doc.add_paragraph()
                # Copy paragraph formatting
                new_p.style = orig_p.style
                new_p.alignment = orig_p.alignment
                new_p.paragraph_format.left_indent = orig_p.paragraph_format.left_indent
                new_p.paragraph_format.right_indent = orig_p.paragraph_format.right_indent
                new_p.paragraph_format.space_before = orig_p.paragraph_format.space_before
                new_p.paragraph_format.space_after = orig_p.paragraph_format.space_after
                new_p.paragraph_format.line_spacing = orig_p.paragraph_format.line_spacing

                # Copy runs with formatting
                for run in orig_p.runs:
                    new_run = new_p.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    if run.font.color.rgb is not None:
                        new_run.font.color.rgb = run.font.color.rgb
                    if run.font.size is not None:
                        new_run.font.size = run.font.size
                    if run.font.name is not None:
                        new_run.font.name = run.font.name

        elif tag == 'delete':
            # Show deleted paragraphs with red strikethrough preserving styles
            for idx in range(i1, i2):
                orig_p = original_doc.paragraphs[idx]
                new_p = doc.add_paragraph()
                # Copy paragraph formatting
                new_p.style = orig_p.style
                new_p.alignment = orig_p.alignment
                new_p.paragraph_format.left_indent = orig_p.paragraph_format.left_indent
                new_p.paragraph_format.right_indent = orig_p.paragraph_format.right_indent
                new_p.paragraph_format.space_before = orig_p.paragraph_format.space_before
                new_p.paragraph_format.space_after = orig_p.paragraph_format.space_after
                new_p.paragraph_format.line_spacing = orig_p.paragraph_format.line_spacing

                for run in orig_p.runs:
                    new_run = new_p.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # Red color
                    new_run.font.strike = True
                    if run.font.size is not None:
                        new_run.font.size = run.font.size
                    if run.font.name is not None:
                        new_run.font.name = run.font.name

        elif tag == 'insert':
            # Show inserted paragraphs with green underline, attempting to inherit formatting
            for para_text in edited_paragraphs[j1:j2]:
                new_p = doc.add_paragraph()
                
                # Determine a reference paragraph to inherit style from
                reference_p = None
                if i1 > 0 and i1 - 1 < len(original_doc.paragraphs):
                    reference_p = original_doc.paragraphs[i1 - 1] # Previous original paragraph
                elif i2 < len(original_doc.paragraphs):
                    reference_p = original_doc.paragraphs[i2] # Next original paragraph (if no previous)
                
                if reference_p:
                    # Inherit paragraph-level formatting
                    new_p.style = reference_p.style
                    new_p.alignment = reference_p.alignment
                    new_p.paragraph_format.left_indent = reference_p.paragraph_format.left_indent
                    new_p.paragraph_format.right_indent = reference_p.paragraph_format.right_indent
                    new_p.paragraph_format.space_before = reference_p.paragraph_format.space_before
                    new_p.paragraph_format.space_after = reference_p.paragraph_format.space_after
                    new_p.paragraph_format.line_spacing = reference_p.paragraph_format.line_spacing
                    
                    # Inherit run-level formatting from the first run of the reference paragraph
                    # This is a heuristic, as the AI's plain text doesn't tell us internal formatting.
                    if reference_p.runs:
                        ref_run = reference_p.runs[0] # Take first run as reference for font properties
                        new_run = new_p.add_run(para_text)
                        new_run.bold = ref_run.bold
                        new_run.italic = ref_run.italic
                        # Set default underline (green for inserted)
                        new_run.font.underline = True 
                        new_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # Green color

                        if ref_run.font.size is not None:
                            new_run.font.size = ref_run.font.size
                        if ref_run.font.name is not None:
                            new_run.font.name = ref_run.font.name
                    else: # Reference paragraph has no runs, or no suitable runs
                        new_run = new_p.add_run(para_text)
                        new_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                        new_run.font.underline = True
                        new_run.font.size = Pt(12) # Default
                        new_run.font.name = 'Calibri' # Default
                else: # No suitable reference paragraph found
                    new_run = new_p.add_run(para_text)
                    new_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                    new_run.font.underline = True
                    new_run.font.size = Pt(12)
                    new_run.font.name = 'Calibri'

        elif tag == 'replace':
            # Show deleted paragraphs in red strikethrough preserving styles
            for idx in range(i1, i2):
                orig_p = original_doc.paragraphs[idx]
                new_p = doc.add_paragraph()
                # Copy paragraph formatting
                new_p.style = orig_p.style
                new_p.alignment = orig_p.alignment
                new_p.paragraph_format.left_indent = orig_p.paragraph_format.left_indent
                new_p.paragraph_format.right_indent = orig_p.paragraph_format.right_indent
                new_p.paragraph_format.space_before = orig_p.paragraph_format.space_before
                new_p.paragraph_format.space_after = orig_p.paragraph_format.space_after
                new_p.paragraph_format.line_spacing = orig_p.paragraph_format.line_spacing

                for run in orig_p.runs:
                    new_run = new_p.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # Red color
                    new_run.font.strike = True
                    if run.font.size is not None:
                        new_run.font.size = run.font.size
                    if run.font.name is not None:
                        new_run.font.name = run.font.name
            
            # Show inserted paragraphs with green underline, attempting to inherit formatting
            for para_text in edited_paragraphs[j1:j2]:
                new_p = doc.add_paragraph()
                
                # Determine a reference paragraph to inherit style from
                reference_p = None
                if i1 > 0 and i1 - 1 < len(original_doc.paragraphs):
                    reference_p = original_doc.paragraphs[i1 - 1] # Previous original paragraph
                elif i2 < len(original_doc.paragraphs):
                    reference_p = original_doc.paragraphs[i2] # Next original paragraph (if no previous)
                
                if reference_p:
                    # Inherit paragraph-level formatting
                    new_p.style = reference_p.style
                    new_p.alignment = reference_p.alignment
                    new_p.paragraph_format.left_indent = reference_p.paragraph_format.left_indent
                    new_p.paragraph_format.right_indent = reference_p.paragraph_format.right_indent
                    new_p.paragraph_format.space_before = reference_p.paragraph_format.space_before
                    new_p.paragraph_format.space_after = reference_p.paragraph_format.space_after
                    new_p.paragraph_format.line_spacing = reference_p.paragraph_format.line_spacing

                    if reference_p.runs:
                        ref_run = reference_p.runs[0] # Take first run as reference
                        new_run = new_p.add_run(para_text)
                        new_run.bold = ref_run.bold
                        new_run.italic = ref_run.italic
                        new_run.font.underline = True 
                        new_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # Green color

                        if ref_run.font.size is not None:
                            new_run.font.size = ref_run.font.size
                        if ref_run.font.name is not None:
                            new_run.font.name = ref_run.font.name
                    else:
                        new_run = new_p.add_run(para_text)
                        new_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                        new_run.font.underline = True
                        new_run.font.size = Pt(12)
                        new_run.font.name = 'Calibri'
                else:
                    new_run = new_p.add_run(para_text)
                    new_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                    new_run.font.underline = True
                    new_run.font.size = Pt(12)
                    new_run.font.name = 'Calibri'
                    
    doc.save(output_path)


def create_side_by_side_diff(original, edited):
    """Create a side-by-side comparison with highlighting"""
    original_lines = original.split('\n')
    edited_lines = edited.split('\n')
    
    diff_html = []
    differ = difflib.SequenceMatcher(None, original_lines, edited_lines)
    
    diff_html.append('<div style="display: flex; font-family: monospace; font-size: 12px;">')
    diff_html.append('<div style="width: 50%; padding-right: 10px; border-right: 1px solid #ccc;">')
    diff_html.append('<h4>Original</h4>')
    
    for tag, i1, i2, j1, j2 in differ.get_opcodes():
        if tag == 'equal':
            for line in original_lines[i1:i2]:
                diff_html.append(f'<div style="padding: 2px;">{line}</div>')
        elif tag == 'delete':
            for line in original_lines[i1:i2]:
                diff_html.append(f'<div style="background-color: #ffcccc; padding: 2px; text-decoration: line-through;">{line}</div>')
        elif tag == 'replace':
            for line in original_lines[i1:i2]:
                diff_html.append(f'<div style="background-color: #ffcccc; padding: 2px; text-decoration: line-through;">{line}</div>')
    
    diff_html.append('</div>')
    diff_html.append('<div style="width: 50%; padding-left: 10px;">')
    diff_html.append('<h4>Edited</h4>')
    
    for tag, i1, i2, j1, j2 in differ.get_opcodes():
        if tag == 'equal':
            for line in edited_lines[j1:j2]:
                diff_html.append(f'<div style="padding: 2px;">{line}</div>')
        elif tag == 'insert':
            for line in edited_lines[j1:j2]:
                diff_html.append(f'<div style="background-color: #ccffcc; padding: 2px; font-weight: bold;">{line}</div>')
        elif tag == 'replace':
            for line in edited_lines[j1:j2]:
                diff_html.append(f'<div style="background-color: #ccffcc; padding: 2px; font-weight: bold;">{line}</div>')
    
    diff_html.append('</div>')
    diff_html.append('</div>')
    
    return ''.join(diff_html)

def chunk_text(text, max_chunk_size=3000):
    """Split text into chunks of max_chunk_size characters without breaking words."""
    words = text.split()
    chunks = []
    current_chunk = []
    current_length = 0
    for word in words:
        if current_length + len(word) + 1 > max_chunk_size:
            chunks.append(' '.join(current_chunk))
            current_chunk = [word]
            current_length = len(word) + 1
        else:
            current_chunk.append(word)
            current_length += len(word) + 1
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    return chunks

def get_edited_text(chat_model, text, selected_rules, custom_instructions=""):
    """Get edited text from AI model with chunking to avoid truncation"""
    try:
        chunks = chunk_text(text)
        edited_chunks = []
        for chunk in chunks:
            prompt = create_editing_prompt(chunk, selected_rules, custom_instructions)
            response = chat_model.invoke([HumanMessage(content=prompt)])
            edited_chunks.append(response.content)
        # Combine all edited chunks
        return '\n'.join(edited_chunks)
    except Exception as e:
        raise Exception(f"Error processing text with AI model: {str(e)}")

def create_diff_view(original, edited):
    """Create a side-by-side diff view"""
    diff = list(difflib.unified_diff(
        original.splitlines(keepends=True),
        edited.splitlines(keepends=True),
        fromfile='Original',
        tofile='Edited',
        n=3
    ))
    return ''.join(diff)

def main():
    st.title("📝 Document Copyediting Tool")
    st.markdown("Upload documents and apply professional copyediting rules using AI")
    
    # Initialize chat model
    chat_model = initialize_chat_model()
    if not chat_model:
        st.stop()
    
    # Sidebar for rule selection
    st.sidebar.header("📋 Copyediting Rules")
    st.sidebar.markdown("Select the rules to apply:")
    
    selected_rules = {}
    for category, rules in STYLE_GUIDE_RULES.items():
        st.sidebar.subheader(category)
        selected_rules[category] = {}
        for rule, default in rules.items():
            selected_rules[category][rule] = st.sidebar.checkbox(
                rule, 
                value=default, 
                key=f"{category}_{rule}"
            )
    
    # Add custom prompt section in sidebar
    st.sidebar.markdown("---")
    st.sidebar.header("✏️ Custom Instructions")
    st.sidebar.markdown("Add any additional editing instructions:")
    
    # Initialize custom_prompt in session state if not already present
    if "custom_prompt" not in st.session_state:
        st.session_state.custom_prompt = ""

    # Add some example prompts
    st.sidebar.markdown("**Example Instructions:**")
    example_prompts = [
        "Make the tone more professional and formal",
        "Convert passive voice to active voice where appropriate",
        "Simplify complex sentences for better readability",
        "Ensure consistent terminology throughout",
        "Add transitional phrases for better flow",
        "Make the writing more concise and direct"
    ]
    
    for i, example in enumerate(example_prompts):
        if st.sidebar.button(f"Use: {example[:30]}...", key=f"example_{i}"):
            st.session_state.custom_prompt = example
            st.rerun() 
            
    # Define the text_area.
    custom_prompt = st.sidebar.text_area(
        "Additional Instructions",
        placeholder="e.g., 'Make the tone more formal', 'Convert to active voice', 'Simplify complex sentences', etc.",
        height=120,
        key="custom_prompt", 
        help="Enter any specific instructions for editing that aren't covered by the standard rules above."
    )
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📤 Upload Document")
        uploaded_file = st.file_uploader(
            "Choose a file",
            type=['txt', 'docx', 'html', 'htm', 'md', 'markdown'],
            help="Supported formats: TXT, DOCX, HTML, Markdown"
        )
        
        if uploaded_file is not None:
            # Display file info
            st.success(f"✅ File uploaded: {uploaded_file.name}")
            st.info(f"File size: {len(uploaded_file.getvalue())} bytes")
            
            # Store the uploaded file's initial position
            uploaded_file.seek(0)
            st.session_state.uploaded_file_bytes = uploaded_file.getvalue()
            
            # Process document (read as text)
            with st.spinner("Processing document..."):
                original_text = process_document(uploaded_file) # This reads the content
                
            if original_text:
                st.subheader("📄 Original Text")
                st.text_area(
                    "Original Content",
                    value=original_text,
                    height=400,
                    key="original_text"
                )
                
                # Show selected editing options
                st.subheader("🔧 Editing Configuration")
                
                any_rule_selected = any(
                    any(rules.values()) for rules in selected_rules.values()
                )
                
                if any_rule_selected:
                    st.write("**Selected Style Rules:**")
                    for category, rules in selected_rules.items():
                        active_rules = [rule for rule, active in rules.items() if active]
                        if active_rules:
                            st.write(f"• **{category}**: {', '.join(active_rules)}")
                else:
                    st.write("*No style guide rules selected*")
                
                # Show custom instructions
                if custom_prompt.strip(): 
                    st.write("**Custom Instructions:**")
                    st.write(f"• {custom_prompt}")
                else:
                    st.write("*No custom instructions provided*")
                
                # Process button
                if st.button("🔄 Apply Copyediting Rules", type="primary"):
                    # Check if any rules are selected or custom prompt is provided
                    if not any_rule_selected and not custom_prompt.strip():
                        st.warning("⚠️ Please select at least one copyediting rule from the sidebar or provide custom instructions.")
                    else:
                        with st.spinner("AI is processing your text..."):
                            try:
                                edited_text = get_edited_text(chat_model, original_text, selected_rules, custom_prompt)
                                st.session_state.edited_text_result = edited_text
                                st.session_state.original_text_result = original_text
                                st.success("✅ Text processed successfully!")
                            except Exception as e:
                                st.error(f"❌ Processing failed: {str(e)}")
    
    with col2:
        st.header("📝 Edited Document")
        
        if hasattr(st.session_state, 'edited_text_result'):
            st.subheader("📄 Edited Text")
            st.text_area(
                "Edited Content",
                value=st.session_state.edited_text_result,
                height=400,
                key="edited_text_display"
            )
            
            # Download options
            st.subheader("💾 Download Options")
            
            col_download1, col_download2 = st.columns(2)
            
            with col_download1:
                # Download edited text
                st.download_button(
                    label="📥 Download Edited Text",
                    data=st.session_state.edited_text_result,
                    file_name=f"edited_{uploaded_file.name if uploaded_file else 'document'}.txt",
                    mime="text/plain"
                )
            
            with col_download2:
                # Download diff
                if hasattr(st.session_state, 'original_text_result'):
                    diff_content = create_diff_view(
                        st.session_state.original_text_result, 
                        st.session_state.edited_text_result
                    )
                    st.download_button(
                        label="📊 Download Diff Report",
                        data=diff_content,
                        file_name=f"diff_{uploaded_file.name if uploaded_file else 'document'}.txt",
                        mime="text/plain"
                    )
        
            # New download button for docx with tracked changes
            if uploaded_file is not None and uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                docx_buffer = io.BytesIO()
                # Use the stored file bytes to recreate a file-like object for docx.Document
                original_docx_bytes_io = io.BytesIO(st.session_state.uploaded_file_bytes)
                format_diff_to_docx(original_docx_bytes_io, st.session_state.edited_text_result, docx_buffer)
                docx_buffer.seek(0)
                st.download_button(
                    label="📄 Download Edited Document (DOCX with tracked changes)",
                    data=docx_buffer.getvalue(),
                    file_name=f"edited_tracked_changes_{uploaded_file.name if uploaded_file else 'document'}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.info("👆 Upload a document and click 'Apply Copyediting Rules' to see the edited version here.")
    
    # Highlighted comparison section (full width)
    if hasattr(st.session_state, 'edited_text_result') and hasattr(st.session_state, 'original_text_result'):
        st.markdown("---")
        st.header("🔍 Highlighted Changes")
        
        # Create tabs for different views
        tab1, tab2, tab3 = st.tabs(["📊 Side-by-Side Comparison", "📝 Highlighted Text", "📋 Diff Report"])
        
        with tab1:
            st.subheader("Side-by-Side Comparison")
            side_by_side_html = create_side_by_side_diff(
                st.session_state.original_text_result, 
                st.session_state.edited_text_result
            )
            st.markdown(side_by_side_html, unsafe_allow_html=True)
        
        with tab2:
            st.subheader("Highlighted Changes")
            col_highlight1, col_highlight2 = st.columns(2)
            
            # Create highlighted versions
            highlighted_original, highlighted_edited = highlight_differences(
                st.session_state.original_text_result, 
                st.session_state.edited_text_result
            )
            
            with col_highlight1:
                st.markdown("**Original (with deletions highlighted)**")
                st.markdown(f'<div style="border: 1px solid #ddd; padding: 10px; height: 400px; overflow-y: scroll; font-family: monospace; font-size: 12px; background-color: #fafafa;">{highlighted_original}</div>', unsafe_allow_html=True)
            
            with col_highlight2:
                st.markdown("**Edited (with additions highlighted)**")
                st.markdown(f'<div style="border: 1px solid #ddd; padding: 10px; height: 400px; overflow-y: scroll; font-family: monospace; font-size: 12px; background-color: #fafafa;">{highlighted_edited}</div>', unsafe_allow_html=True)
        
        with tab3:
            st.subheader("Traditional Diff Report")
            diff_content = create_diff_view(
                st.session_state.original_text_result, 
                st.session_state.edited_text_result
            )
            st.code(diff_content, language="diff")
        
        # Legend
        st.markdown("---")
        st.subheader("🎨 Legend")
        st.markdown("""
        <div style="display: flex; gap: 20px; margin: 10px 0;">
            <div style="display: flex; align-items: center; gap: 5px;">
                <span style="background-color: #ffcccc; padding: 2px 5px; text-decoration: line-through;">Deleted Text</span>
                <span>- Text removed from original</span>
            </div>
            <div style="display: flex; align-items: center; gap: 5px;">
                <span style="background-color: #ccffcc; padding: 2px 5px; font-weight: bold;">Added Text</span>
                <span>- Text added in edited version</span>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    # Instructions
    st.markdown("---")
    st.subheader("📖 How to Use")
    st.markdown("""
    1. **Select Rules**: Choose the copyediting rules you want to apply from the sidebar
    2. **Add Custom Instructions**: Optionally, provide specific editing instructions in the text area
    3. **Upload Document**: Upload a supported file (TXT, DOCX, HTML, Markdown)
    4. **Process**: Click "Apply Copyediting Rules" to generate the edited version
    5. **Review**: Compare the original and edited text using the different view options
    6. **Download**: Save the edited document or diff report
    """)

if __name__ == '__main__':
    main()
